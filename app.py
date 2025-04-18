import streamlit as st
import requests
from PIL import Image
from io import BytesIO
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import zipfile

st.set_page_config(page_title="Imagens para Word", layout="wide")
st.title("📎 Ferramenta Spot - VExpenses")

MAX_ZIP_SIZE_MB = 150
MAX_ZIP_SIZE_BYTES = MAX_ZIP_SIZE_MB * 1024 * 1024

def extrair_links_e_ids(file):
    wb = load_workbook(file, data_only=True)
    ws = wb.active
    headers = {cell.value: idx for idx, cell in enumerate(ws[1])}

    if "Link do Anexo" not in headers or "ID da Despesa" not in headers or "ID do Relatório" not in headers:
        raise ValueError("A planilha deve conter as colunas 'Link do Anexo', 'ID da Despesa' e 'ID do Relatório'.")

    col_link = headers["Link do Anexo"]
    col_id_despesa = headers["ID da Despesa"]
    col_id_relatorio = headers["ID do Relatório"]

    dados = []
    for row in ws.iter_rows(min_row=2):
        linha_excel = row[0].row
        id_despesa = row[col_id_despesa].value
        id_relatorio = row[col_id_relatorio].value
        cell_link = row[col_link]
        url = cell_link.hyperlink.target if cell_link.hyperlink else None
        dados.append((linha_excel, id_despesa, id_relatorio, url))

    return dados

def pdf_para_imagens(pdf_bytes):
    imagens = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            img = Image.open(BytesIO(pix.tobytes("png")))
            imagens.append(img)
    return imagens

def ajustar_altura_doc_paragrafo(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))
    pPr.append(OxmlElement('w:keepNext'))

def inserir_imagem_redimensionada(paragraph, img, largura_max=5.5, altura_max=7):
    img_io = BytesIO()
    img.save(img_io, format='PNG')
    img_io.seek(0)
    largura, altura = img.size
    escala = min((largura_max * 96) / largura, (altura_max * 96) / altura)
    escala *= 1.1
    nova_largura = largura * escala / 96
    run = paragraph.add_run()
    run.add_picture(img_io, width=Inches(nova_largura))

def aplicar_fonte_arial(run):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)

uploaded_file = st.file_uploader("📂 Envie a planilha (.xlsx) com os links de imagem", type=["xlsx"])
manual_uploads = {}

if uploaded_file:
    try:
        info_links = extrair_links_e_ids(uploaded_file)

        if not info_links:
            st.error("❌ Nenhum link encontrado na planilha.")
        else:
            st.success(f"✅ {len(info_links)} registros encontrados.")
            st.header("📸 Upload manual para anexos ausentes")

            for linha, id_despesa, id_relatorio, url in info_links:
                if not url:
                    st.warning(f"🔍 Imagem ausente: ID da Despesa {id_despesa} | ID do Relatório {id_relatorio}")
                    img = st.file_uploader(
                        f"Envie a imagem para linha {linha}",
                        type=["jpg", "png", "jpeg"],
                        key=f"upload_linha_{linha}_despesa_{id_despesa}"
                    )
                    manual_uploads[linha] = img

            imagens_pendentes = [linha for linha, _, _, url in info_links if not url and not manual_uploads.get(linha)]
            if imagens_pendentes:
                st.info("⏳ Aguardando envio de todas as imagens manuais antes de gerar o Word.")
                st.stop()

            if st.button("📝 Gerar Documento Word"):
                erros = []
                doc = Document()
                log_area = st.empty()
                imagens_com_nomes = []

                for i, (linha, id_despesa, id_relatorio, url) in enumerate(info_links, 1):
                    try:
                        log_area.markdown(f"🔄 Processando linha **{linha}**: ID da Despesa `{id_despesa}` / ID do Relatório `{id_relatorio}`")

                        if not url:
                            uploaded_image = manual_uploads.get(linha)
                            if not uploaded_image:
                                raise ValueError("Imagem manual não enviada.")
                            img = Image.open(uploaded_image).convert("RGB")
                            imagens = [img]
                        else:
                            if not url.startswith("http"):
                                url = "https://" + url
                            response = requests.get(url, timeout=20)
                            response.raise_for_status()
                            content_type = response.headers.get('Content-Type', '')

                            if 'pdf' in content_type:
                                imagens = pdf_para_imagens(response.content)
                            else:
                                img = Image.open(BytesIO(response.content)).convert("RGB")
                                extrema = img.getextrema()
                                if all(e[0] == e[1] for e in extrema):
                                    raise ValueError("Imagem aparentemente em branco.")
                                imagens = [img]

                        for idx, img in enumerate(imagens):
                            nome_img = f"{id_despesa}_pag{idx+1}.png" if len(imagens) > 1 else f"{id_despesa}.png"
                            imagens_com_nomes.append((nome_img, img))
                            doc.add_page_break()
                            p = doc.add_paragraph()
                            ajustar_altura_doc_paragrafo(p)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run(f"ID da Despesa: {id_despesa} / ID do Relatório: {id_relatorio}")
                            aplicar_fonte_arial(run)
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            inserir_imagem_redimensionada(p_img, img)

                    except Exception as e:
                        erros.append((linha, id_despesa, id_relatorio, e))
                        doc.add_page_break()
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(f"ID da Despesa: {id_despesa} / ID do Relatório: {id_relatorio}")
                        aplicar_fonte_arial(run)
                        p_err = doc.add_paragraph()
                        p_err.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run_err = p_err.add_run("⚠️ Erro ao carregar imagem: " + str(e))
                        aplicar_fonte_arial(run_err)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                log_area.empty()

                st.success("✅ Documento Word gerado com sucesso!")
                st.download_button(
                    label="📥 Baixar Word",
                    data=buffer,
                    file_name="anexos_ordenados.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.markdown("### 📦 Download de imagens zipadas (máx. 150MB cada)")
                pacote_atual = BytesIO()
                zip_atual = zipfile.ZipFile(pacote_atual, mode="w", compression=zipfile.ZIP_DEFLATED)
                tamanho_atual = 0
                pacote_idx = 1
                lista_buffers = []

                for nome, imagem in imagens_com_nomes:
                    buffer_img = BytesIO()
                    imagem.save(buffer_img, format="PNG")
                    tamanho_img = buffer_img.tell()
                    buffer_img.seek(0)

                    if tamanho_atual + tamanho_img > MAX_ZIP_SIZE_BYTES:
                        zip_atual.close()
                        pacote_atual.seek(0)
                        lista_buffers.append((pacote_idx, pacote_atual))
                        pacote_idx += 1
                        pacote_atual = BytesIO()
                        zip_atual = zipfile.ZipFile(pacote_atual, mode="w", compression=zipfile.ZIP_DEFLATED)
                        tamanho_atual = 0

                    zip_atual.writestr(nome, buffer_img.read())
                    tamanho_atual += tamanho_img

                zip_atual.close()
                pacote_atual.seek(0)
                lista_buffers.append((pacote_idx, pacote_atual))

                for idx, zip_buffer in lista_buffers:
                    st.download_button(
                        label=f"📁 Baixar pacote ZIP {idx}",
                        data=zip_buffer,
                        file_name=f"imagens_pacote_{idx}.zip",
                        mime="application/zip"
                    )

                if erros:
                    st.markdown("### ❌ Falhas detectadas")
                    for linha, id_despesa, id_relatorio, erro in erros:
                        st.write(f"Linha {linha} | Despesa: {id_despesa} | Relatório: {id_relatorio} → {erro}")

    except Exception as e:
        st.error(f"Erro ao processar: {e}")
